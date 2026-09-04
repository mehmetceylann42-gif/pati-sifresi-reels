#!/usr/bin/env python3
"""Storyboard'dan seslendirme senaryosu üretir: Word (.docx) + OKU.txt.

    python scripts/make_script_doc.py 29-hayvana-eziyet-sorusturma-sarti
    python scripts/make_script_doc.py <slug> --out "C:/Users/efeka/Desktop"
    python scripts/make_script_doc.py <slug> --only alternatif      # tek beat

Üretilenler:
    <out>/<slug>-senaryo.docx              okunacak metin, beat beat
    <out>/<slug>-<rol>-yeniden-kayit.docx  --only verildiğinde
    voice_recordings/<slug>/OKU.txt        dosya adı → cümle eşlemesi

`--only`, bir beat'in metni yayından sonra değiştiğinde işe yarar: kalan
kayıtlar geçerli kaldığı için Word'e yalnızca yeniden okunacak beat konur ve
dosya adı hangi WAV'ın üzerine yazılacağını birebir söyler.

Kullanıcı kendi sesiyle kaydettiğinde her beat ayrı bir WAV olur ve
storyboard'daki `voice_file` alanı o dosyayı gösterir; build_reel.py
kaydı sentez sesin yerine koyar. Bu script iki tarafın da aynı dosya
adlarını kullanmasını garanti eder — elle yazılınca kaçan şey buydu.
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

PROJECT_ROOT = Path(__file__).resolve().parent.parent
STORYBOARDS = PROJECT_ROOT / "content" / "storyboards.json"
WORDS_PER_SECOND = 2.1  # 18 videoluk build_log kalibrasyonu (2026-09-03)

AMBER = RGBColor(0xB0, 0x73, 0x0E)
SLATE = RGBColor(0x4A, 0x5A, 0x6A)


def load(slug: str) -> dict:
    for entry in json.loads(STORYBOARDS.read_text(encoding="utf-8")):
        if entry.get("slug") == slug:
            return entry
    raise SystemExit(f"'{slug}' content/storyboards.json içinde yok.")


def beat_file(board: dict, index: int, beat: dict) -> str:
    declared = beat.get("voice_file")
    if declared:
        return Path(declared).name
    return f"beat_{index:02d}_{beat.get('role', 'beat')}.wav"


def write_oku(board: dict) -> Path:
    slug = board["slug"]
    folder = PROJECT_ROOT / "voice_recordings" / slug
    folder.mkdir(parents=True, exist_ok=True)
    beats = board["beats"]

    lines = [
        f"{slug} — seslendirme kaydı",
        "",
        f"Bu klasore su {len(beats)} WAV dosyasi konacak (isimler birebir):",
        "",
    ]
    for index, beat in enumerate(beats):
        lines.append(beat_file(board, index, beat))
        lines.append(f"  {beat['vo']}")
        lines.append("")

    path = folder / "OKU.txt"
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def para(document, text, *, size=11, bold=False, color=None, space_after=6,
         italic=False, align=None):
    paragraph = document.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    return paragraph


def select_beats(board: dict, only: list[str] | None) -> list[int]:
    """--only ile verilen rol adlarını / 0 tabanlı indeksleri beat sırasına çevirir."""
    if not only:
        return list(range(len(board["beats"])))
    picked: list[int] = []
    for token in only:
        token = token.strip()
        match = [i for i, b in enumerate(board["beats"]) if b.get("role") == token]
        if not match and token.isdigit() and int(token) < len(board["beats"]):
            match = [int(token)]
        if not match:
            roles = ", ".join(b.get("role", "?") for b in board["beats"])
            raise SystemExit(f"'{token}' bu storyboard'da yok. Roller: {roles}")
        picked.extend(match)
    return sorted(set(picked))


def write_docx(board: dict, out_dir: Path, only: list[str] | None = None) -> Path:
    slug = board["slug"]
    beats = board["beats"]
    picked = select_beats(board, only)
    partial = len(picked) < len(beats)
    words = sum(len(b["vo"].split()) for b in beats)

    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Slug ASCII'dir; başlıkta Türkçe yazım korunmalı (marka yazım kuralı),
    # o yüzden varsa storyboard'daki `title` alanı kullanılır.
    para(document, board.get("title") or slug.replace("-", " ").upper(),
         size=18, bold=True, space_after=2)
    para(document, "Seslendirme senaryosu · PATİ ŞİFRESİ", size=12, color=AMBER,
         bold=True, space_after=14)

    meta = [
        ("Slug", slug),
        ("Kaynak", board.get("source_name", "")),
        ("Kaynak linki", board.get("source_url", "")),
        ("İddia sınıfı", f"{board.get('claim_class', '?')} · doğrulama: {board.get('verified_on', '—')} · yetki alanı: {board.get('jurisdiction', '—')}"),
        ("Beat sayısı", f"{len(beats)} · {words} kelime · tahmini {words / WORDS_PER_SECOND:.0f} saniye"),
        ("Müzik", board.get("music_id", "—")),
    ]
    for label, value in meta:
        if not value:
            continue
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        head = paragraph.add_run(f"{label}: ")
        head.bold = True
        head.font.size = Pt(9.5)
        body = paragraph.add_run(str(value))
        body.font.size = Pt(9.5)
        body.font.color.rgb = SLATE

    if board.get("scope_note"):
        document.add_paragraph()
        para(document, "KAPSAM CÜMLESİ (K4) — bu sınır videoda söylenir",
             size=9.5, bold=True, color=AMBER, space_after=3)
        para(document, board["scope_note"], size=9.5, italic=True, color=SLATE, space_after=12)

    document.add_paragraph()
    para(document, "YENİDEN OKUNACAK METİN" if partial else "OKUNACAK METİN",
         size=13, bold=True, space_after=4)
    if partial:
        para(document,
             f"Bu dosyada {len(beats)} beat'in yalnızca {len(picked)} tanesi var — "
             "metni değişen beat(ler). Diğer kayıtlara dokunma; aşağıdaki WAV "
             "dosyasının ÜZERİNE yaz, dosya adını değiştirme.",
             size=9.5, italic=True, color=AMBER, space_after=12)
    else:
        para(document, "Her başlık ayrı bir WAV dosyasıdır. Dosya adlarını birebir koru.",
             size=9.5, italic=True, color=SLATE, space_after=12)

    for index in picked:
        beat = beats[index]
        filename = beat_file(board, index, beat)
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(f"{index + 1}.  {filename}")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = AMBER

        for chunk in beat.get("chunks") or [beat["vo"]]:
            para(document, chunk, size=14, space_after=2)

        notes = []
        if beat.get("emph"):
            notes.append("vurgu: " + ", ".join(beat["emph"]))
        if beat.get("pitch"):
            notes.append(f"ton: {beat['pitch']}")
        if beat.get("pause_after"):
            notes.append(f"sonrasında {beat['pause_after']} sn duraklama")
        if notes:
            para(document, " · ".join(notes), size=9, italic=True, color=SLATE, space_after=2)

    document.add_page_break()
    para(document, "Seslendirme notları", size=13, bold=True, space_after=8)
    notes = [
        "Her beat'i AYRI dosya olarak kaydet; dosya adları yukarıdaki gibi olmalı.",
        "Kuru ve ölçülü oku — abartılı tonlama profesyonellikten uzaklaştırıyor.",
        "Büyük harfle yazılan kelimeler vurgulanacak yerlerdir.",
        "Beat sonundaki duraklamaları kayda dahil etme; kurgu zaten ekliyor.",
        "Cümlenin başında ve sonunda yarım saniye sessizlik bırak, kırpma kolaylaşsın.",
        f"Toplam hedef süre: yaklaşık {words / WORDS_PER_SECOND:.0f} saniye "
        "(sabit üst sınır yok — bkz. CONTENT_AND_COMMERCE_RULES.md süre kuralı).",
    ]
    # Bazı videolarda tonlama içeriğin kendisidir; genel "kuru oku" notu orada
    # yanlış yönlendirir. Storyboard kendi notunu ekleyebilir.
    notes.extend(board.get("voice_notes") or [])
    for note in notes:
        paragraph = document.add_paragraph(note, style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        for run in paragraph.runs:
            run.font.size = Pt(10.5)

    para(document, "", space_after=8)
    para(document, "Kayıtlar nereye konacak", size=13, bold=True, space_after=4)
    para(document, f"Hayvan-Kanali/voice_recordings/{slug}/", size=10.5, space_after=10)
    para(document, "Sonra:  python scripts/build_reel.py " + slug, size=10.5, space_after=4)

    out_dir.mkdir(parents=True, exist_ok=True)
    if partial:
        tag = "-".join(beats[i].get("role", str(i)) for i in picked)
        path = out_dir / f"{slug}-{tag}-yeniden-kayit.docx"
    else:
        path = out_dir / f"{slug}-senaryo.docx"
    document.save(str(path))
    return path


def main() -> int:
    parser = argparse.ArgumentParser(description="Storyboard'dan .docx senaryo + OKU.txt üretir")
    parser.add_argument("slug")
    parser.add_argument("--out", default=str(Path.home() / "Desktop"),
                        help="Word dosyasının yazılacağı klasör (varsayılan: Masaüstü)")
    parser.add_argument("--only", nargs="+", metavar="ROL",
                        help="Yalnızca bu beat(ler) için Word üret (rol adı ya da sıra numarası)")
    args = parser.parse_args()

    board = load(args.slug)
    oku = write_oku(board)
    doc = write_docx(board, Path(args.out), args.only)
    print(f"Word  : {doc}")
    print(f"OKU   : {oku.relative_to(PROJECT_ROOT)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
